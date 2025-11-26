import json
from docx import Document
import os
from tempfile import NamedTemporaryFile


def validate_assessments_weight(assessments):
    total = sum([int(a['weight']) for a in assessments])
    return total == 100


def export_to_docx(outline):
    doc = Document()

    doc.add_heading(f"{outline.course_code} – {outline.course_title}", level=1)

    doc.add_heading("Course Information", level=2)
    doc.add_paragraph(f"Term: {outline.term}")
    doc.add_paragraph(f"Credits: {outline.credits}")
    doc.add_paragraph(f"Modality: {outline.modality}")
    doc.add_paragraph(f"Language: {outline.language}")

    doc.add_heading("Instructor", level=2)
    doc.add_paragraph(f"Name: {outline.professor_name}")
    doc.add_paragraph(f"Email: {outline.professor_email}")
    doc.add_paragraph(f"Office Hours: {outline.office_hours}")
    doc.add_paragraph(f"Location: {outline.location}")

    doc.add_heading("Assessments", level=2)
    for a in json.loads(outline.assessments_json or "[]"):
        doc.add_paragraph(f"- {a['name']} ({a['weight']}%)")

    doc.add_heading("Schedule", level=2)
    for w in json.loads(outline.schedule_json or "[]"):
        doc.add_paragraph(f"Week {w['week']}: {w['topic']}")

    doc.add_heading("Policies", level=2)
    doc.add_paragraph(outline.policies_text or "")

    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
